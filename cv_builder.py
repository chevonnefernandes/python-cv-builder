from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


speak(
    "Welcome to your personal CV generator. Make sure your CV headshot is saved as headshot.png"
)

document = Document()

# profile picture
document.add_picture("headshot.png", width=Inches(2.0))

# contact details
name = input("Enter your full name? ")
speak(f"Hi {name}. Next, let's add your email address")
email = input("Enter your email address? ")
speak("Great. Now enter your LinkedIn url")
linkedin_url = input("Enter your LinkedIn url: ")

document.add_paragraph(
    f"{name.capitalize()} | {email.lower()} | {linkedin_url.lower()}"
)

# professional profile
speak(
    "Let's add a professional profile to concisely summarise your skills, experiences and career goals."
)
document.add_heading("Professional Profile")
about_me = input("Add your professional profile here: ")
document.add_paragraph(about_me)

# technical skills
speak(
    "Let's add some technical skills to highlight specific technical abilities and competencies."
)
document.add_heading("Technical Skills")
skill = input("Enter technical skill: ")
p = document.add_paragraph(skill)
p.style = "List Bullet"

# additional technical skills
while True:
    add_skill = input("Do you have an additional skill you'd like to add? Yes or No: ")
    if add_skill.lower() == "yes":
        skill = input("Enter technical skill: ")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break

# work experience
speak("Now let's add in your relevant work experience.")
document.add_heading("Experience")
p = document.add_paragraph()

company = input("Company name: ")
start_date = input('Start date (in the format "Month YYYY"): ')
end_date = input('End date (in the format "Month YYYY"): ')
experience_details = input(f"Describe your role at {company}: ")

p.add_run(f"{company} ").bold = True
p.add_run(f"{start_date} - {end_date} \n").italic = True
p.add_run(experience_details)

# additional work experience
while True:
    add_exp = input(
        "Do you have an additional experience you would like to add? Yes or No: "
    )
    if add_exp.lower() == "yes":
        p = document.add_paragraph()

        company = input("Company name: ")
        start_date = input('Start date (in the format "Month YYYY"): ')
        end_date = input('End date (in the format "Month YYYY"): ')
        experience_details = input(f"Describe your role at {company}: ")

        p.add_run(f"{company} ").bold = True
        p.add_run(f"{start_date} - {end_date} \n").italic = True
        p.add_run(experience_details)
    else:
        break

# footer
section = document.sections[0].footer
p = section.paragraphs[0]
p.text = "References available on request."

speak(f"Great job {name} Your CV has been generated.")

document.save("cv.docx")
